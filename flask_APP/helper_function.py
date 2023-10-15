import os
import pandas as pd
import ast
import openai
import re
from typing import Tuple, List
import concurrent.futures
from tqdm import tqdm
from scipy import spatial
from docx import Document as DocxDocument
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
import concurrent.futures
import logging
import time



from transformers import GPT2Tokenizer
tokenizer = GPT2Tokenizer.from_pretrained("gpt2")
# Please refer to OpenAI website for embedding model and GPT model selection
# https://platform.openai.com/docs/introduction
EMBEDDING_MODEL = 'text-embedding-ada-002'
EMBEDDING_CTX_LENGTH = 8191
EMBEDDING_ENCODING = 'cl100k_base'
embeddings_buffer_tokens = 7500 
GPT_MODEL = 'gpt-3.5-turbo-16k' 
TOKEN_BUDGET = 16384 - 500 

# utility function 

def count_tokens(text):
    tokens = tokenizer.tokenize(text)
    return len(tokens)

def truncate_text(text, max_tokens):
    tokens = tokenizer.tokenize(text)
    return ' '.join(tokens[:max_tokens])




# Most related secton search
tfidf_vectorizer = TfidfVectorizer(stop_words='english')

def combined_similarity(raw_text1, raw_text2, emb1, emb2, alpha=0.7):
    # Embedding similarity
    emb_sim = 1 - spatial.distance.cosine(emb1, emb2)
    
    # TF-IDF similarity
    tfidf_matrix = tfidf_vectorizer.fit_transform([raw_text1, raw_text2])
    tfidf_sim = cosine_similarity(tfidf_matrix[0:1], tfidf_matrix)[0][1]
    
    # Combine similarities
    combined_sim = alpha * emb_sim + (1 - alpha) * tfidf_sim
    return combined_sim

def strings_ranked_by_relatedness(query, df, relatedness_fn=combined_similarity, top_n=4):
    query_embedding_response = openai.Embedding.create(
        model=EMBEDDING_MODEL,
        input=query,
    )
    query_embedding = query_embedding_response["data"][0]["embedding"]
    
    strings_sections_and_relatednesses = [
        (row["section"], row["text"], relatedness_fn(query, row["text"], query_embedding, row["embeddings"]))
        for i, row in df.iterrows()
    ]

    strings_sections_and_relatednesses.sort(key=lambda x: x[2], reverse=True)
    sections, strings, relatednesses = zip(*strings_sections_and_relatednesses)
    return sections[:top_n], strings[:top_n], relatednesses[:top_n]

def calculate_embeddings_section_text_df(df, model='text-embedding-ada-002', max_tokens=EMBEDDING_CTX_LENGTH):
    embeddings = []

    for text in df['text']:
        # Check if the text exceeds the maximum token limit
        tokens_count = count_tokens(text)
        if tokens_count > max_tokens:
            print(f"Warning: The text with {tokens_count} tokens exceeds the maximum limit of {max_tokens} tokens. It will be truncated, meaning information will be lost.")
            text = truncate_text(text, max_tokens)

        # Create embedding with OpenAI's Embedding API
        response = openai.Embedding.create(
            input=text,
            model=model
        )
        
        chunk_embeddings = response['data'][0]['embedding']
        embeddings.append(chunk_embeddings)

    # Add the embeddings as a new column to the DataFrame
    df['embeddings'] = embeddings

    return df



def string_to_list(s):
    """Converts a string representation of a list to an actual list."""
    if isinstance(s, list):
        return s
    elif isinstance(s, str):
        try:
            return ast.literal_eval(s)
        except Exception as e:
            print(f"Error occurred: {str(e)}")
            return None
    else:
        print(f"Unexpected input type: {type(s)}, expected string or list.")
        return None
    
def inspect_styles(doc):
    styles = set()
    for para in doc.paragraphs:
        styles.add(para.style.name)
    return styles

def save_to_parquet(df, path, filename):
    # Create the directories in the path if they don't exist
    os.makedirs(path, exist_ok=True)
    
    # Combine the path and filename to create the full file path
    full_path = os.path.join(path, filename)
    
    # Save the DataFrame to the Parquet file at the specified path
    df.to_parquet(full_path)
    print(f"File saved to {full_path}")
    
    return full_path  # Return the full path

def load_from_parquet(path, filename):
    # Combine the path and filename to create the full file path
    full_path = os.path.join(path, filename)
    
    # Load the DataFrame from the Parquet file at the specified path
    return pd.read_parquet(full_path)

def save_selected_columns_to_csv(df, columns, path, filename):
    # Create the directories in the path if they don't exist
    os.makedirs(path, exist_ok=True)
    
    # Combine the path and filename to create the full file path
    full_path = os.path.join(path, filename)
    
    # Select only the specified columns
    selected_df = df[columns]
    
    # Save the selected columns to a CSV file at the specified path
    selected_df.to_csv(full_path, index=False)
    print(f"File saved to {full_path}")

    return full_path



# process documents
def split_into_sections_with_GPT_Filtering_progress_bar(filepath, embedding_max_tokens=embeddings_buffer_tokens, word_limit_for_paragraph_content=600, max_tokens=TOKEN_BUDGET):
    doc = DocxDocument(filepath)
    filename = os.path.basename(filepath)
    styles_used = inspect_styles(doc)
    
    # Estimate total work (modify this according to your needs)
    total_work = len(doc.paragraphs)
    
    
    # Filter the styles to get the heading styles
    heading_styles = {style for style in styles_used if 'Heading' in style or 'SubTitle' in style}

    sections = []
    section_title = ""
    section_content = []
    
    with tqdm(total=total_work, desc="Processing document") as main_progress_bar:    
        for para_index, para in enumerate(doc.paragraphs):
            main_progress_bar.set_description(f"Processing paragraph {para_index + 1}/{total_work}")
            # time.sleep(1)
            # main_progress_bar.update(1)
            if para.style.name in heading_styles:  # if the paragraph is a heading
                if section_content:  # if there's any content accumulated
                    content = '\n'.join(section_content)
                    if content.strip() != "[]" and (count_tokens(content) > embedding_max_tokens or len(content.split()) > word_limit_for_paragraph_content):
                        print(f"Warning: section '{section_title}' contains more than {embedding_max_tokens} tokens or {word_limit_for_paragraph_content} words.\nContents in that section will be filtered by {GPT_MODEL}")
                        
                        # Check if the content exceeds the max token limit for the model
                        if count_tokens(content) > max_tokens:
                            print(f"Warning: section '{section_title}' contains more than {max_tokens} tokens for {GPT_MODEL} to filter.\nThe content will be truncated.")
                            content = truncate_text(content, max_tokens)
                        
                        print(f"Filtering using {GPT_MODEL} ")
                        additional_sections = parse_and_split_sections_by_GPT(document_text=content)
                        print("Completed filtering")
                        for additional_section in additional_sections:
                            # additional_section['section'] = additional_section['section']
                            # additional_section['text'] = f"[{additional_section['section']}]\n{additional_section['text']}"
                            # additional_section['section'] = f"[{filename}]\n{additional_section['section']}"
                            additional_section_text = additional_section['section']
                            additional_section['text'] = f"[{additional_section_text}]\n{additional_section['text']}"
                            additional_section['section'] = f"[{filename}]\n{additional_section_text}"
                        
                        sections.extend(additional_sections)
                    elif content.strip() != "[]":
                        if content.strip() == f'[{section_title}]':
                            content += " Section header with no specific content"
                        sections.append({
                            "section": section_title,
                            "text": content
                        })
                    section_content = []

                section_title_text = para.text.rstrip() # removing trailing whitespaces here
                section_content.append(f"[{section_title_text}]")
                section_title = f"[{filename}]\n{section_title_text}" # Add filename here
            else:
                section_content.append(para.text)
            main_progress_bar.update(1)
        
        if section_content:  # append the last section if any
            content = '\n'.join(section_content)
            if content.strip() != "[]" and (count_tokens(content) > embedding_max_tokens or len(content.split()) > word_limit_for_paragraph_content):
                print(f"Warning: section '{section_title}' contains more than {embedding_max_tokens} tokens or {word_limit_for_paragraph_content} words.")
                
                # Check if the content exceeds the max token limit for the model
                if count_tokens(content) > max_tokens:
                    print(f"Warning: section '{section_title}' contains more than {max_tokens} tokens. The content will be truncated.")
                    content = truncate_text(content, max_tokens)
                
                print("Processing Last Section")
                additional_sections = parse_and_split_sections_by_GPT(content)
                
                for additional_section in additional_sections:
                    # additional_section['text'] = f"[{additional_section['section']}]\n{additional_section['text']}"
                    # additional_section['section'] = f"[{filename}]\n{additional_section['section']}"
                    additional_section_text = additional_section['section']
                    additional_section['text'] = f"[{additional_section_text}]\n{additional_section['text']}"
                    additional_section['section'] = f"[{filename}]\n{additional_section_text}"

                sections.extend(additional_sections)
            elif content.strip() != "[]":
                sections.append({
                    "section": section_title,
                    "text": content
                })
        

    # Close the progress bar when done
    main_progress_bar.close()
    return sections

def parse_and_split_sections_by_GPT(
        document_text : str,
        system_message = """You will be given a text extracted from a document. This text might contain multiple sections separated by titles. 
        Each section starts with a title, typically a line on its own, followed by its content. Your task is to parse this document and identify these sections. 
        Please return a list of dictionaries, with each dictionary representing a section in the following format: 
        {
            "section": section_title,
            "text": content
        }. 
        If the document only contains a single section with no distinct section title, represent it as a single dictionary with the section title as 'Main'.""" ,
        system_template = """Please identify the section titles and their corresponding text content from the given document. 
        Section titles usually appear as standalone lines preceding the section content. 
        Format your response as a list of dictionaries, each representing a section in the format: 
        {
            "section": section_title,
            "text": content
        }. 
        If the document only contains a single section with no distinct section title, represent it as a single dictionary with the section title as 'Main'."""    
    ):

    if count_tokens(document_text) > TOKEN_BUDGET:
        print(f"Text token exceeded limit, your text length is {count_tokens(document_text)} where the max token for {GPT_MODEL} is {TOKEN_BUDGET}")
        return 
    
    messages = [
        {"role": "system", "content": system_message},
        {"role": "system", "content": system_template},
        {"role": "user", "content": document_text},
    ]
    response = openai.ChatCompletion.create(
        model=GPT_MODEL,
        messages= messages,
        temperature=0
    )
    return string_to_list(response.choices[0].message['content'])

def prepare_document(
    filepath : str,
    verbose: bool = False,
    
):
    try:
        # section_text_df = pd.DataFrame(split_into_sections_with_GPT_Filtering_progress_bar_Test(filepath=filepath))
        section_text_df = pd.DataFrame(split_into_sections_with_GPT_Filtering_progress_bar(filepath=filepath))
        if verbose:
            print("Finished spliting documents!")
    except Exception as e:
        print("Error extracting text for sections:", e)
        return str(e)

    try:
        if verbose:
            print("Calculating embeddings!")
        embedding_df = calculate_embeddings_section_text_df(df=section_text_df, max_tokens=EMBEDDING_CTX_LENGTH)
        if verbose:
            print("Done calculating embeddings!")
    except Exception as e:
        print("Error calculating embeddings for section text dataframe:", e)
        return str(e)
    
    return embedding_df



# ask quesiton section
def create_response_dataframe(sections: list[str], strings: list[str], relatedness_scores: list[float]) -> pd.DataFrame:
    """Create a DataFrame with three columns, 'Section', 'Relatedness Score' and 'Data'."""
    data = {
        'Section Name': sections,
        'Relatedness Score %': [f"{round(score * 100)} %" for score in relatedness_scores],
        'Data': strings,
    }
    df = pd.DataFrame(data)
    
    # Modify 'Data' column for rows that match the pattern [content], where content can be any word, space or common punctuation
    pattern = r"^\[[a-zA-Z0-9_ \-\.,;]+\]$"
    additional_description = " This is a section title with no content. Please refer to the document by searching the section name!"
    df['Data'] = df['Data'].apply(lambda x: x + additional_description if re.match(pattern, x) else x)
    
    return df

def merge_dataframes(dfs: List[pd.DataFrame]) -> pd.DataFrame:
    """
    Accepts a list of dataframes and merges them together if they have the same columns.
    """
    # Check if all dataframes have the same columns
    if not all(df.columns.equals(dfs[0].columns) for df in dfs):
        raise ValueError("All dataframes should have the same columns.")
    
    # Merge dataframes
    merged_df = pd.concat(dfs, ignore_index=True)
    return merged_df

def ask_question_mutiple_doc_and_return_df_response_drop_duplicate(question: str, dfs: List[pd.DataFrame], top_n: int = 5) -> pd.DataFrame:
    """
    Takes a list of dataframes as input, checks if all dataframes have the same columns,
    then merges those dataframes together and uses "strings_ranked_by_relatedness" 
    to find the most related section based on the provided question.
    """
    merged_df = merge_dataframes(dfs)
    sections, strings, relatedness_scores = strings_ranked_by_relatedness(query=question, df=merged_df, top_n=top_n * 2)  # get a bit more than needed to account for possible duplicates
    
    # Convert to DataFrame early to drop duplicates
    preliminary_df = pd.DataFrame({
        'Section Name': sections,
        'Relatedness Score %': [f"{round(score * 100)} %" for score in relatedness_scores],
        'Data': strings
    })
    
    # Drop duplicates
    preliminary_df.drop_duplicates(subset=['Section Name', 'Data'], inplace=True)
    
    # Truncate to top_n rows
    truncated_df = preliminary_df.head(top_n)
    
    df_response = create_response_dataframe(truncated_df['Section Name'].tolist(), truncated_df['Data'].tolist(), truncated_df['Relatedness Score %'].str.rstrip(' %').astype(float) / 100.0)
    df_response = df_response.set_index([pd.Index([question]*df_response.shape[0]), df_response.index])
    df_response.index.names = ['user_question', 'original_index']

    return df_response


# global handbook code

def create_global_handbook(base_df: pd.DataFrame, dfs_list: list[pd.DataFrame], target_names: list[str] = ["UK", "Ireland"]) -> list[pd.DataFrame]:
    """
    Uses the base_df (US handbook) as a source of questions and asks them to each dataframe in dfs_list.
    Returns a list of resulting DataFrames.
    """
    results = []

    for idx, df_with_embedding_2 in enumerate(dfs_list):
        target_name = target_names[idx]
        result_data = {
            f'US Handbook Section': [],
            f'US Handbook Section Text': [],
            f'{target_name} Handbook Related Section': [],
            f'{target_name} Handbook Related Text': [],
            f'US {target_name} Handbook Relatedness Score %': []
        }

        for _, row in base_df.iterrows():
            question = row['text']
            response_df = ask_question_mutiple_doc_and_return_df_response_drop_duplicate(question, df_with_embedding_2, top_n=1)
            most_relevant_row = response_df.iloc[0]  # get the most relevant row

            result_data[f'US Handbook Section'].append(row['section'])
            result_data[f'US Handbook Section Text'].append(question)
            result_data[f'{target_name} Handbook Related Section'].append(most_relevant_row['Section'])
            result_data[f'{target_name} Handbook Related Text'].append(most_relevant_row['Data'])
            result_data[f'US {target_name} Handbook Relatedness Score %'].append(most_relevant_row['Relatedness Score %'])

        results.append(pd.DataFrame(result_data))

    return results

def string_to_list(string):
    # Note: You haven't provided the implementation for this function.
    # Assuming it converts a string to a list in some way.
    return string.split("\n")  # Placeholder implementation


system_message_US_base = """
You are given sections from handbooks representing the U.S., UK, and Ireland respectively. The U.S. section serves as the blueprint. 
Your task is to synthesize the contents, focusing on commonalities while highlighting significant contrasts among the regions. 
The final output should be coherent, not exceed twice the length of the U.S. section, and should maintain the primary structure of the U.S. blueprint.
"""

system_template_US_base = """
Given the U.S. blueprint section provided, blend in details from the UK and Ireland sections.
Summarize the content while highlighting common themes and contrasting points from all three regions. 
Use indicators like "In the US", "In the UK", etc., for clarity. Ensure the output is concise and structured like the U.S. section.
"""

system_message_all_equal = """
You are presented with sections from handbooks representing the U.S., UK, and Ireland. 
Your task is to synthesize the content from all three regions into a concise and coherent section. 
The goal is to provide a summary that captures the essence of each region while highlighting any significant contrasts among them.
"""

system_template_all_equal = """
Given the sections from the U.S., UK, and Ireland, blend the details to produce a concise summary.
Highlight common themes and emphasize contrasting points from each region using indicators like "In the US", "In the UK", etc. 
Ensure the output is brief, clear, and provides a comprehensive view of all three sections.
"""


def US_base_merge_paragraphs(
        blueprint: str, 
        *reference_strings: str,
        system_message=system_message_US_base,
        system_template=system_template_US_base
    ):
    
    # Construct the messages for the model
    labeled_content = "US blueprint section:\n" + blueprint
    if reference_strings:
        labeled_content += "\n\nUK Relevant section:\n" + reference_strings[0]
        if len(reference_strings) > 1:
            labeled_content += "\n\nIreland Relevant section:\n" + reference_strings[1]
    
    messages = [
        {"role": "system", "content": system_message},
        {"role": "system", "content": system_template},
        {"role": "user", "content": labeled_content}
    ]

    # Call the GPT model
    response = openai.ChatCompletion.create(
        model=GPT_MODEL,
        messages=messages,
        temperature=0
    )

    # Post-process the response
    merged_text = response.choices[0].message['content']
    # Remove square brackets and double quotes
    cleaned_text = merged_text.replace("[", "").replace("]", "").replace('"', '')
    
    return cleaned_text

def All_HandBook_merge_paragraphs(
        blueprint: str, 
        *reference_strings: str,
        system_message=system_message_all_equal,
        system_template=system_template_all_equal
    ):
    
    # Construct the messages for the model
    labeled_content = "US blueprint section:\n" + blueprint
    if reference_strings:
        labeled_content += "\n\nUK Relevant section:\n" + reference_strings[0]
        if len(reference_strings) > 1:
            labeled_content += "\n\nIreland Relevant section:\n" + reference_strings[1]
    
    messages = [
        {"role": "system", "content": system_message},
        {"role": "system", "content": system_template},
        {"role": "user", "content": labeled_content}
    ]

    # Call the GPT model
    response = openai.ChatCompletion.create(
        model=GPT_MODEL,
        messages=messages,
        temperature=0
    )

    # Post-process the response
    merged_text = response.choices[0].message['content']
    # Remove square brackets and double quotes
    cleaned_text = merged_text.replace("[", "").replace("]", "").replace('"', '')
    
    return cleaned_text


logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

MAX_RETRIES = 5
RATE_LIMIT_SLEEP_DURATION = 10  # 10 seconds, adjust if needed

def process_section(row, dfs_list):
    question = row['text']
    uk_relatedness = strings_ranked_by_relatedness(question, dfs_list[0], top_n=1)[2][0]
    ireland_relatedness = strings_ranked_by_relatedness(question, dfs_list[1], top_n=1)[2][0]

    if uk_relatedness >= 0.65 or ireland_relatedness >= 0.65:
        uk_text = strings_ranked_by_relatedness(question, dfs_list[0], top_n=1)[1][0]
        ireland_text = strings_ranked_by_relatedness(question, dfs_list[1], top_n=1)[1][0]
        
        retries = 0
        while retries < MAX_RETRIES:
            try:
                merged_text = US_base_merge_paragraphs(question, uk_text, ireland_text)
                 # Remove any square brackets from the merged_text
                merged_text = merged_text.replace("[", "").replace("]", "")
                # Determine the status based on relatedness scores
                if uk_relatedness >= 0.65 and ireland_relatedness >= 0.65:
                    status = "Merged with UK and Ireland"
                elif uk_relatedness >= 0.65:
                    status = "Merged with UK"
                elif ireland_relatedness >= 0.65:
                    status = "Merged with Ireland"
                else:
                    status = "Error: Unexpected merging state"
                
                return (row['section'], question, merged_text, status)
                
            except openai.error.OpenAIError as e:
                if e.status == 429:  # Rate limit exceeded
                    logger.warning(f"Rate limit hit. Sleeping for {RATE_LIMIT_SLEEP_DURATION} seconds...")
                    time.sleep(RATE_LIMIT_SLEEP_DURATION)
                    retries += 1
                    continue
                else:
                    logger.error(f"Error while processing section: {e}. Retrying...")
                    retries += 1
            except Exception as e:
                logger.error(f"Unknown error: {e}. Retrying...")
                retries += 1
        logger.error(f"Failed to process section after {MAX_RETRIES} retries. Skipping...")
        return (row['section'], question, None, "Error")
    else:
        return (row['section'], question, question, "Unchanged")

def process_section_Mixed(row, dfs_list):
    question = row['text']
    uk_relatedness = strings_ranked_by_relatedness(question, dfs_list[0], top_n=1)[2][0]
    ireland_relatedness = strings_ranked_by_relatedness(question, dfs_list[1], top_n=1)[2][0]

    if uk_relatedness >= 0.65 or ireland_relatedness >= 0.65:
        uk_text = strings_ranked_by_relatedness(question, dfs_list[0], top_n=1)[1][0]
        ireland_text = strings_ranked_by_relatedness(question, dfs_list[1], top_n=1)[1][0]
        
        retries = 0
        while retries < MAX_RETRIES:
            try:
                merged_text = All_HandBook_merge_paragraphs(question, uk_text, ireland_text)
                 # Remove any square brackets from the merged_text
                merged_text = merged_text.replace("[", "").replace("]", "")
                # Determine the status based on relatedness scores
                if uk_relatedness >= 0.65 and ireland_relatedness >= 0.65:
                    status = "Merged with UK and Ireland"
                elif uk_relatedness >= 0.65:
                    status = "Merged with UK"
                elif ireland_relatedness >= 0.65:
                    status = "Merged with Ireland"
                else:
                    status = "Error: Unexpected merging state"
                
                return (row['section'], question, merged_text, status)
                
            except openai.error.OpenAIError as e:
                if e.status == 429:  # Rate limit exceeded
                    logger.warning(f"Rate limit hit. Sleeping for {RATE_LIMIT_SLEEP_DURATION} seconds...")
                    time.sleep(RATE_LIMIT_SLEEP_DURATION)
                    retries += 1
                    continue
                else:
                    logger.error(f"Error while processing section: {e}. Retrying...")
                    retries += 1
            except Exception as e:
                logger.error(f"Unknown error: {e}. Retrying...")
                retries += 1
        logger.error(f"Failed to process section after {MAX_RETRIES} retries. Skipping...")
        return (row['section'], question, None, "Error")
    else:
        return (row['section'], question, question, "Unchanged")


def create_combined_global_handbook_US_Based(base_df: pd.DataFrame, dfs_list: list[pd.DataFrame], excel_filename : str) -> pd.DataFrame:
    result_data = {
        'US Handbook Section': [],
        'US Handbook Section Text': [],
        'Global Handbook Text': [],
        'Status': []
    }

    with concurrent.futures.ThreadPoolExecutor() as executor:
        # Store each future with its original position/index
        futures = {executor.submit(process_section, row, dfs_list): idx for idx, (_, row) in enumerate(base_df.iterrows())}
        temp_results = []  # Temporary storage for results with their index
        
        for future in tqdm(concurrent.futures.as_completed(futures), total=len(futures), desc=f"Processing sections"):
            idx = futures[future]
            try:
                result = future.result()
                temp_results.append((idx, result))
            except Exception as e:
                logger.error(f"Error processing a section: {e}")
                result = (None, None, None, "Error")
                temp_results.append((idx, result))

        # Sort results by their index to preserve order
        temp_results.sort(key=lambda x: x[0])
        ordered_results = [res for _, res in temp_results]

        # Populate the result_data with ordered results
        for section, question, merged_text, status in ordered_results:
            section = section.replace("[", "").replace("]", "") if section else section
            question = question.replace("[", "").replace("]", "") if question else question
            merged_text = merged_text.replace("[", "").replace("]", "") if merged_text else merged_text
            result_data['US Handbook Section'].append(section)
            result_data['US Handbook Section Text'].append(question)
            if merged_text:
                result_data['Global Handbook Text'].append(merged_text)
            result_data['Status'].append(status)

    result_df = pd.DataFrame(result_data)
    result_df.to_excel(excel_filename, index=False)
    return result_df

def create_combined_global_handbook_Mixed(base_df: pd.DataFrame, dfs_list: list[pd.DataFrame], excel_filename : str) -> pd.DataFrame:
    result_data = {
        'US Handbook Section': [],
        'US Handbook Section Text': [],
        'Global Handbook Text': [],
        'Status': []
    }

    with concurrent.futures.ThreadPoolExecutor() as executor:
        # Store each future with its original position/index
        futures = {executor.submit(process_section_Mixed, row, dfs_list): idx for idx, (_, row) in enumerate(base_df.iterrows())}
        temp_results = []  # Temporary storage for results with their index
        
        for future in tqdm(concurrent.futures.as_completed(futures), total=len(futures), desc=f"Processing sections"):
            idx = futures[future]
            try:
                result = future.result()
                temp_results.append((idx, result))
            except Exception as e:
                logger.error(f"Error processing a section: {e}")
                result = (None, None, None, "Error")
                temp_results.append((idx, result))

        # Sort results by their index to preserve order
        temp_results.sort(key=lambda x: x[0])
        ordered_results = [res for _, res in temp_results]

        # Populate the result_data with ordered results
        for section, question, merged_text, status in ordered_results:
            section = section.replace("[", "").replace("]", "") if section else section
            question = question.replace("[", "").replace("]", "") if question else question
            merged_text = merged_text.replace("[", "").replace("]", "") if merged_text else merged_text
            result_data['US Handbook Section'].append(section)
            result_data['US Handbook Section Text'].append(question)
            if merged_text:
                result_data['Global Handbook Text'].append(merged_text)
            result_data['Status'].append(status)

    result_df = pd.DataFrame(result_data)
    result_df.to_excel(excel_filename, index=False)
    return result_df

